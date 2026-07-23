from pathlib import Path
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from sqlalchemy import text
import re


BASE_DIR = Path(__file__).resolve().parents[1]
TEMPLATE_PRIMER_LLAMADO = BASE_DIR / "templates" / "rrll" / "abandono" / "primer_llamado_abandono.docx"
TIPO_DOC_PRIMER_LLAMADO = 13

TEMPLATE_SEGUNDO_LLAMADO = BASE_DIR / "templates" / "rrll" / "abandono" / "segundo_llamado_abandono.docx"
TIPO_DOC_SEGUNDO_LLAMADO = 14

TEMPLATE_CARTA_FINALIZACION = BASE_DIR / "templates" / "rrll" / "finalizacion" / "carta_finalizacion_contrato.docx"
TIPO_DOC_CARTA_FINALIZACION = 4

TEMPLATE_PAQUETE_RETIRO = BASE_DIR / "templates" / "rrll" / "paquete" / "paquete_retiro.docx"
TIPO_DOC_PAQUETE_RETIRO = 10

TEMPLATE_PAQUETE_RETIRO_VOLUNTARIO = BASE_DIR / "templates" / "rrll" / "paquete" / "paquete_retiro_voluntario.docx"

FIRMA_YENY = BASE_DIR / "assets" / "comunicaciones" / "FIRMA_YENY.png"

# ✅ RUTA CORRECTA:
# cada documento generado debe quedar dentro de:
# app/storage/rrll/retiros/{IdRetiroLaboral}/archivo.docx
OUTPUT_BASE_DIR = Path("C:/LaPerfeccionStorage/rrll/retiros")


def _clean_text(value):
    """
    Limpia tabs, saltos de línea y espacios múltiples para que Word
    no distribuya el texto raro al reemplazar placeholders.
    """
    if value is None:
        return ""

    text_value = str(value)
    text_value = text_value.replace("\t", " ")
    text_value = text_value.replace("\r", " ")
    text_value = text_value.replace("\n", " ")
    text_value = re.sub(r"\s+", " ", text_value)

    return text_value.strip()


def _upper_text(value):
    if value is None:
        return ""
    return str(value).upper().strip()


def _replace_text_in_paragraph(paragraph, replacements: dict):
    for key, value in replacements.items():
        if key in paragraph.text:
            for run in paragraph.runs:
                if key in run.text:
                    run.text = run.text.replace(key, str(value))


def _replace_text_in_table(table, replacements: dict):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_text_in_paragraph(paragraph, replacements)



def _paragraph_has_image(paragraph) -> bool:
    """Indica si el párrafo contiene una imagen incrustada."""
    return bool(
        paragraph._p.xpath(".//w:drawing")
        or paragraph._p.xpath(".//w:pict")
    )


def _remove_images_from_paragraph(paragraph):
    """Elimina únicamente las imágenes contenidas en el párrafo indicado."""
    for drawing in paragraph._p.xpath(".//w:drawing"):
        parent = drawing.getparent()
        if parent is not None:
            parent.remove(drawing)

    for pict in paragraph._p.xpath(".//w:pict"):
        parent = pict.getparent()
        if parent is not None:
            parent.remove(pict)


def _clear_paragraph(paragraph):
    """Limpia el contenido del párrafo conservando su posición en la plantilla."""
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def _iter_document_paragraph_groups(doc):
    """
    Devuelve grupos de párrafos del cuerpo y de las tablas para poder
    encontrar la firma sin modificar logos, encabezados u otras imágenes.
    """
    yield doc.paragraphs

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield cell.paragraphs


def _insertar_firma_yeny(doc, modo_compacto: bool = False):
    """
    Reemplaza los placeholders de nombre/cargo por la imagen completa de Yeny.
    Si existe una firma vieja inmediatamente antes de los placeholders,
    elimina solamente esa imagen y conserva el resto de la plantilla.
    """
    if not FIRMA_YENY.exists():
        raise FileNotFoundError(f"No se encontró la firma de Yeny: {FIRMA_YENY}")

    firma_insertada = False

    for paragraphs in _iter_document_paragraph_groups(doc):
        for index, paragraph in enumerate(paragraphs):
            if "{{NOMBRE_ANALISTA}}" not in paragraph.text:
                continue

            # Busca hacia atrás únicamente la imagen más cercana a la firma.
            for previous_index in range(index - 1, max(-1, index - 4), -1):
                previous_paragraph = paragraphs[previous_index]
                if _paragraph_has_image(previous_paragraph):
                    _remove_images_from_paragraph(previous_paragraph)
                    break

            _clear_paragraph(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            if modo_compacto:
                # Solo para la carta de finalización: evita que Word envíe
                # todo el bloque de firma a una segunda hoja.
                paragraph.paragraph_format.keep_together = False
                paragraph.paragraph_format.keep_with_next = False
                paragraph.paragraph_format.space_before = 0
                paragraph.paragraph_format.space_after = 0
                ancho_firma = Inches(2.10)
            else:
                # Se conserva exactamente el formato que ya quedó validado
                # para primer llamado, segundo llamado y paquetes de retiro.
                paragraph.paragraph_format.keep_together = True
                paragraph.paragraph_format.keep_with_next = True
                ancho_firma = Inches(2.35)

            run = paragraph.add_run()
            run.add_picture(str(FIRMA_YENY), width=ancho_firma)

            texto_firma = paragraph.add_run(
                "\nYENY CUESTO"
                "\nANALISTA TALENTO HUMANO"
                "\nAseos La Perfección S.A.S."
            )
            texto_firma.bold = True

            firma_insertada = True

            # El cargo ya viene dentro de la nueva imagen.
            if index + 1 < len(paragraphs):
                next_paragraph = paragraphs[index + 1]
                if "{{CARGO_ANALISTA}}" in next_paragraph.text:
                    _clear_paragraph(next_paragraph)

                    # En algunas plantillas aparece también el nombre de la empresa
                    # debajo del cargo. Se retira para dejar únicamente el bloque
                    # completo de la firma nueva.
                    if index + 2 < len(paragraphs):
                        company_paragraph = paragraphs[index + 2]
                        if company_paragraph.text.strip().upper() == "ASEOS LA PERFECCIÓN S.A.S.":
                            _clear_paragraph(company_paragraph)

    if not firma_insertada:
        raise ValueError(
            "No se encontró el placeholder {{NOMBRE_ANALISTA}} "
            "en la plantilla seleccionada."
        )

def _get_output_dir_for_retiro(id_retiro_laboral: int) -> Path:
    output_dir = OUTPUT_BASE_DIR / str(id_retiro_laboral)
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir


def _build_output_path(id_retiro_laboral: int, prefix: str) -> Path:
    output_dir = _get_output_dir_for_retiro(id_retiro_laboral)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return output_dir / f"{prefix}_{id_retiro_laboral}_{timestamp}.docx"


def obtener_datos_primer_llamado(db, id_retiro_laboral: int):
    query = text("""
        SELECT
            rl."IdRetiroLaboral",
            rp."NumeroIdentificacion" AS "NumeroDocumento",
            TRIM(
                COALESCE(rp."Nombres", '') || ' ' ||
                COALESCE(rp."Apellidos", '')
            ) AS "NombreCompleto",
            COALESCE(da."Direccion", '') AS "Direccion",
            COALESCE(da."Barrio", '') AS "Barrio",
            COALESCE(rp."Celular", '') AS "Telefono",
            COALESCE(ca."NombreCargo", '') AS "Cargo",
            COALESCE(psy."FechaUltimoDiaLaborado", rl."FechaRetiro") AS "FechaAusencia"
        FROM public."RetiroLaboral" rl
        INNER JOIN public."RegistroPersonal" rp
            ON rl."IdRegistroPersonal" = rp."IdRegistroPersonal"
        LEFT JOIN public."DatosAdicionales" da
            ON rp."IdRegistroPersonal" = da."IdRegistroPersonal"
        LEFT JOIN public."AsignacionCargoCliente" acc
            ON acc."IdRegistroPersonal" = rp."IdRegistroPersonal"
        LEFT JOIN public."Cargo" ca
            ON acc."IdCargo" = ca."IdCargo"
        LEFT JOIN public."PazYSalvoOperaciones" psy
            ON psy."IdRetiroLaboral" = rl."IdRetiroLaboral"
        WHERE rl."IdRetiroLaboral" = :id_retiro_laboral
        LIMIT 1;
    """)

    row = db.execute(query, {"id_retiro_laboral": id_retiro_laboral}).mappings().first()

    if not row:
        raise ValueError(f"No se encontraron datos para IdRetiroLaboral={id_retiro_laboral}")

    datos = dict(row)
    datos["NumeroDocumento"] = _clean_text(datos.get("NumeroDocumento"))
    datos["NombreCompleto"] = _clean_text(datos.get("NombreCompleto"))
    datos["Direccion"] = _clean_text(datos.get("Direccion"))
    datos["Barrio"] = _clean_text(datos.get("Barrio"))
    datos["Telefono"] = _clean_text(datos.get("Telefono"))
    datos["Cargo"] = _clean_text(datos.get("Cargo"))

    return datos


def generar_primer_llamado(db, id_retiro_laboral: int):
    if not TEMPLATE_PRIMER_LLAMADO.exists():
        raise FileNotFoundError(f"No se encontró la plantilla: {TEMPLATE_PRIMER_LLAMADO}")

    datos = obtener_datos_primer_llamado(db, id_retiro_laboral)
    doc = Document(str(TEMPLATE_PRIMER_LLAMADO))
    _insertar_firma_yeny(doc)

    fecha_ausencia = datos.get("FechaAusencia")
    if fecha_ausencia:
        try:
            fecha_ausencia = fecha_ausencia.strftime("%d/%m/%Y")
        except Exception:
            fecha_ausencia = _clean_text(fecha_ausencia)
    else:
        fecha_ausencia = ""

    replacements = {
        "{{FECHA_HOY}}": datetime.today().strftime("%d/%m/%Y"),
        "{{NOMBRE_COMPLETO}}": _upper_text(datos.get("NombreCompleto", "")),
        "{{NUMERO_DOCUMENTO}}": _upper_text(datos.get("NumeroDocumento", "")),
        "{{DIRECCION}}": _upper_text(datos.get("Direccion", "")),
        "{{BARRIO}}": _upper_text(datos.get("Barrio", "")),
        "{{TELEFONO}}": _upper_text(datos.get("Telefono", "")),
        "{{CARGO}}": _upper_text(datos.get("Cargo", "")),
        "{{CIUDAD}}": "CIUDAD",
        "{{FECHA_AUSENCIA}}": fecha_ausencia,
        "{{ASUNTO}}": "PRIMER LLAMADO ABANDONO INASISTENCIA AL CARGO",
    }

    for paragraph in doc.paragraphs:
        _replace_text_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        _replace_text_in_table(table, replacements)

    output_path = _build_output_path(id_retiro_laboral, "primer_llamado_retiro")
    doc.save(str(output_path))
    return output_path


def generar_segundo_llamado(db, id_retiro_laboral: int):
    if not TEMPLATE_SEGUNDO_LLAMADO.exists():
        raise FileNotFoundError(f"No se encontró la plantilla: {TEMPLATE_SEGUNDO_LLAMADO}")

    datos = obtener_datos_primer_llamado(db, id_retiro_laboral)
    doc = Document(str(TEMPLATE_SEGUNDO_LLAMADO))
    _insertar_firma_yeny(doc)

    fecha_ausencia = datos.get("FechaAusencia")
    if fecha_ausencia:
        try:
            fecha_ausencia = fecha_ausencia.strftime("%d/%m/%Y")
        except Exception:
            fecha_ausencia = _clean_text(fecha_ausencia)
    else:
        fecha_ausencia = ""

    replacements = {
        "{{FECHA_HOY}}": datetime.today().strftime("%d/%m/%Y"),
        "{{NOMBRE_COMPLETO}}": _upper_text(datos.get("NombreCompleto", "")),
        "{{NUMERO_DOCUMENTO}}": _upper_text(datos.get("NumeroDocumento", "")),
        "{{DIRECCION}}": _upper_text(datos.get("Direccion", "")),
        "{{BARRIO}}": _upper_text(datos.get("Barrio", "")),
        "{{TELEFONO}}": _upper_text(datos.get("Telefono", "")),
        "{{CARGO}}": _upper_text(datos.get("Cargo", "")),
        "{{CIUDAD}}": "CIUDAD",
        "{{FECHA_AUSENCIA}}": fecha_ausencia,
        "{{ASUNTO}}": "SEGUNDO LLAMADO ABANDONO INASISTENCIA AL CARGO",
    }

    for paragraph in doc.paragraphs:
        _replace_text_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        _replace_text_in_table(table, replacements)

    output_path = _build_output_path(id_retiro_laboral, "segundo_llamado_retiro")
    doc.save(str(output_path))
    return output_path


def generar_carta_finalizacion(db, id_retiro_laboral: int):
    if not TEMPLATE_CARTA_FINALIZACION.exists():
        raise FileNotFoundError(f"No se encontró la plantilla: {TEMPLATE_CARTA_FINALIZACION}")

    datos = obtener_datos_primer_llamado(db, id_retiro_laboral)
    doc = Document(str(TEMPLATE_CARTA_FINALIZACION))
    _insertar_firma_yeny(doc, modo_compacto=True)

    fecha_ausencia = datos.get("FechaAusencia")
    if fecha_ausencia:
        try:
            fecha_ausencia = fecha_ausencia.strftime("%d/%m/%Y")
        except Exception:
            fecha_ausencia = _clean_text(fecha_ausencia)
    else:
        fecha_ausencia = ""

    replacements = {
        "{{FECHA_HOY}}": datetime.today().strftime("%d/%m/%Y"),
        "{{NOMBRE_COMPLETO}}": _upper_text(datos.get("NombreCompleto", "")),
        "{{NUMERO_DOCUMENTO}}": _upper_text(datos.get("NumeroDocumento", "")),
        "{{DIRECCION}}": _upper_text(datos.get("Direccion", "")),
        "{{BARRIO}}": _upper_text(datos.get("Barrio", "")),
        "{{TELEFONO}}": _upper_text(datos.get("Telefono", "")),
        "{{CARGO}}": _upper_text(datos.get("Cargo", "")),
        "{{CIUDAD}}": "CIUDAD",
        "{{FECHA_AUSENCIA}}": fecha_ausencia,
        "{{ASUNTO}}": "CARTA DE FINALIZACIÓN DEL CONTRATO",
    }

    for paragraph in doc.paragraphs:
        _replace_text_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        _replace_text_in_table(table, replacements)

    output_path = _build_output_path(id_retiro_laboral, "carta_finalizacion_retiro")
    doc.save(str(output_path))
    return output_path


def generar_paquete_retiro(db, id_retiro_laboral: int):
    if not TEMPLATE_PAQUETE_RETIRO.exists():
        raise FileNotFoundError(f"No se encontró la plantilla: {TEMPLATE_PAQUETE_RETIRO}")

    datos = obtener_datos_primer_llamado(db, id_retiro_laboral)
    print("DEBUG DATOS PAQUETE:", datos)

    q_motivo = text("""
        SELECT "IdMotivoRetiro"
        FROM public."RetiroLaboral"
        WHERE "IdRetiroLaboral" = :id_retiro_laboral
        LIMIT 1;
    """)

    row_motivo = db.execute(q_motivo, {
        "id_retiro_laboral": id_retiro_laboral
    }).mappings().first()

    id_motivo = row_motivo["IdMotivoRetiro"] if row_motivo else None
    print("DEBUG ID MOTIVO RETIRO:", id_motivo)

    es_voluntario = (id_motivo == 1)

    if es_voluntario:
        print("DEBUG PAQUETE: usando plantilla VOLUNTARIO")
        doc = Document(str(TEMPLATE_PAQUETE_RETIRO_VOLUNTARIO))
    else:
        print("DEBUG PAQUETE: usando plantilla NORMAL")
        doc = Document(str(TEMPLATE_PAQUETE_RETIRO))

    _insertar_firma_yeny(doc)

    fecha_fin = datos.get("FechaAusencia")
    if fecha_fin:
        try:
            fecha_fin = fecha_fin.strftime("%d/%m/%Y")
        except Exception:
            fecha_fin = str(fecha_fin)
    else:
        fecha_fin = ""

    replacements = {
        "{{FECHA_HOY}}": datetime.today().strftime("%d/%m/%Y"),
        "{{FECHA_FIN}}": fecha_fin,
        "{{NOMBRE_COMPLETO}}": _upper_text(datos.get("NombreCompleto", "")),
        "{{NUMERO_DOCUMENTO}}": _upper_text(datos.get("NumeroDocumento", "")),
        "{{DIRECCION}}": _upper_text(datos.get("Direccion", "")),
        "{{BARRIO}}": _upper_text(datos.get("Barrio", "")),
        "{{TELEFONO}}": _upper_text(datos.get("Telefono", "")),
        "{{CARGO}}": _upper_text(datos.get("Cargo", "")),
        "{{CIUDAD}}": "CIUDAD",
    }

    for paragraph in doc.paragraphs:
        _replace_text_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        _replace_text_in_table(table, replacements)

    output_path = _build_output_path(id_retiro_laboral, "paquete_retiro")
    doc.save(str(output_path))
    return output_path


def generar_y_registrar_primer_llamado(
    db,
    id_retiro_laboral: int,
    usuario_actualizacion: str = "RRLL"
):
    output_path = generar_primer_llamado(db, id_retiro_laboral)

    if not output_path.exists():
        raise FileNotFoundError("No se pudo generar físicamente el documento.")

    nombre_archivo = output_path.name
    nombre_original = output_path.name
    ruta_archivo = str(output_path).replace("\\", "/")
    extension_archivo = output_path.suffix.lower()
    peso_archivo = output_path.stat().st_size
    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    q_old = text("""
        SELECT
            "IdRetiroLaboralAdjunto",
            "RutaArchivo"
        FROM public."RetiroLaboralAdjunto"
        WHERE "IdRetiroLaboral" = :id_retiro_laboral
          AND "IdTipoDocumentoRetiro" = :id_tipo_documento_retiro
          AND COALESCE("Eliminado", false) = false
          AND COALESCE("Activo", true) = true
        ORDER BY "IdRetiroLaboralAdjunto" DESC
        LIMIT 1;
    """)

    old = db.execute(q_old, {
        "id_retiro_laboral": id_retiro_laboral,
        "id_tipo_documento_retiro": TIPO_DOC_PRIMER_LLAMADO
    }).mappings().first()

    if old:
        q_desactivar = text("""
            UPDATE public."RetiroLaboralAdjunto"
            SET
                "Activo" = false,
                "Eliminado" = true,
                "FechaActualizacion" = now(),
                "UsuarioActualizacion" = :usuario_actualizacion
            WHERE "IdRetiroLaboralAdjunto" = :id_adjunto;
        """)
        db.execute(q_desactivar, {
            "id_adjunto": old["IdRetiroLaboralAdjunto"],
            "usuario_actualizacion": usuario_actualizacion
        })

        ruta_old = Path(old["RutaArchivo"]) if old.get("RutaArchivo") else None
        if ruta_old and ruta_old.exists():
            ruta_old.unlink(missing_ok=True)

    q_insert = text("""
        INSERT INTO public."RetiroLaboralAdjunto" (
            "IdRetiroLaboral",
            "IdTipoDocumentoRetiro",
            "NombreArchivo",
            "NombreArchivoOriginal",
            "RutaArchivo",
            "ExtensionArchivo",
            "PesoArchivo",
            "Observacion",
            "OrigenArchivo",
            "MimeType",
            "Activo",
            "Eliminado",
            "FechaCreacion",
            "FechaActualizacion",
            "CreadoPor",
            "UsuarioActualizacion"
        )
        VALUES (
            :id_retiro_laboral,
            :id_tipo_documento_retiro,
            :nombre_archivo,
            :nombre_archivo_original,
            :ruta_archivo,
            :extension_archivo,
            :peso_archivo,
            :observacion,
            'GENERADO',
            :mime_type,
            true,
            false,
            now(),
            now(),
            :creado_por,
            :usuario_actualizacion
        )
        RETURNING
            "IdRetiroLaboralAdjunto",
            "IdRetiroLaboral",
            "IdTipoDocumentoRetiro",
            "NombreArchivo",
            "NombreArchivoOriginal",
            "RutaArchivo",
            "ExtensionArchivo",
            "PesoArchivo",
            "Observacion",
            "OrigenArchivo",
            "MimeType",
            "Activo";
    """)

    row = db.execute(q_insert, {
        "id_retiro_laboral": id_retiro_laboral,
        "id_tipo_documento_retiro": TIPO_DOC_PRIMER_LLAMADO,
        "nombre_archivo": nombre_archivo,
        "nombre_archivo_original": nombre_original,
        "ruta_archivo": ruta_archivo,
        "extension_archivo": extension_archivo,
        "peso_archivo": peso_archivo,
        "observacion": "Documento generado automáticamente: Primer llamado",
        "mime_type": mime_type,
        "creado_por": usuario_actualizacion,
        "usuario_actualizacion": usuario_actualizacion,
    }).mappings().first()

    db.commit()
    return dict(row)


def generar_y_registrar_segundo_llamado(
    db,
    id_retiro_laboral: int,
    usuario_actualizacion: str = "RRLL"
):
    output_path = generar_segundo_llamado(db, id_retiro_laboral)

    if not output_path.exists():
        raise FileNotFoundError("No se pudo generar físicamente el documento.")

    nombre_archivo = output_path.name
    nombre_original = output_path.name
    ruta_archivo = str(output_path).replace("\\", "/")
    extension_archivo = output_path.suffix.lower()
    peso_archivo = output_path.stat().st_size
    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    q_old = text("""
        SELECT
            "IdRetiroLaboralAdjunto",
            "RutaArchivo"
        FROM public."RetiroLaboralAdjunto"
        WHERE "IdRetiroLaboral" = :id_retiro_laboral
          AND "IdTipoDocumentoRetiro" = :id_tipo_documento_retiro
          AND COALESCE("Eliminado", false) = false
          AND COALESCE("Activo", true) = true
        ORDER BY "IdRetiroLaboralAdjunto" DESC
        LIMIT 1;
    """)

    old = db.execute(q_old, {
        "id_retiro_laboral": id_retiro_laboral,
        "id_tipo_documento_retiro": TIPO_DOC_SEGUNDO_LLAMADO
    }).mappings().first()

    if old:
        q_desactivar = text("""
            UPDATE public."RetiroLaboralAdjunto"
            SET
                "Activo" = false,
                "Eliminado" = true,
                "FechaActualizacion" = now(),
                "UsuarioActualizacion" = :usuario_actualizacion
            WHERE "IdRetiroLaboralAdjunto" = :id_adjunto;
        """)
        db.execute(q_desactivar, {
            "id_adjunto": old["IdRetiroLaboralAdjunto"],
            "usuario_actualizacion": usuario_actualizacion
        })

        ruta_old = Path(old["RutaArchivo"]) if old.get("RutaArchivo") else None
        if ruta_old and ruta_old.exists():
            ruta_old.unlink(missing_ok=True)

    q_insert = text("""
        INSERT INTO public."RetiroLaboralAdjunto" (
            "IdRetiroLaboral",
            "IdTipoDocumentoRetiro",
            "NombreArchivo",
            "NombreArchivoOriginal",
            "RutaArchivo",
            "ExtensionArchivo",
            "PesoArchivo",
            "Observacion",
            "OrigenArchivo",
            "MimeType",
            "Activo",
            "Eliminado",
            "FechaCreacion",
            "FechaActualizacion",
            "CreadoPor",
            "UsuarioActualizacion"
        )
        VALUES (
            :id_retiro_laboral,
            :id_tipo_documento_retiro,
            :nombre_archivo,
            :nombre_archivo_original,
            :ruta_archivo,
            :extension_archivo,
            :peso_archivo,
            :observacion,
            'GENERADO',
            :mime_type,
            true,
            false,
            now(),
            now(),
            :creado_por,
            :usuario_actualizacion
        )
        RETURNING
            "IdRetiroLaboralAdjunto",
            "IdRetiroLaboral",
            "IdTipoDocumentoRetiro",
            "NombreArchivo",
            "NombreArchivoOriginal",
            "RutaArchivo",
            "ExtensionArchivo",
            "PesoArchivo",
            "Observacion",
            "OrigenArchivo",
            "MimeType",
            "Activo";
    """)

    row = db.execute(q_insert, {
        "id_retiro_laboral": id_retiro_laboral,
        "id_tipo_documento_retiro": TIPO_DOC_SEGUNDO_LLAMADO,
        "nombre_archivo": nombre_archivo,
        "nombre_archivo_original": nombre_original,
        "ruta_archivo": ruta_archivo,
        "extension_archivo": extension_archivo,
        "peso_archivo": peso_archivo,
        "observacion": "Documento generado automáticamente: Segundo llamado",
        "mime_type": mime_type,
        "creado_por": usuario_actualizacion,
        "usuario_actualizacion": usuario_actualizacion,
    }).mappings().first()

    db.commit()
    return dict(row)


def generar_y_registrar_carta_finalizacion(
    db,
    id_retiro_laboral: int,
    usuario_actualizacion: str = "RRLL"
):
    output_path = generar_carta_finalizacion(db, id_retiro_laboral)

    if not output_path.exists():
        raise FileNotFoundError("No se pudo generar físicamente el documento.")

    nombre_archivo = output_path.name
    nombre_original = output_path.name
    ruta_archivo = str(output_path).replace("\\", "/")
    extension_archivo = output_path.suffix.lower()
    peso_archivo = output_path.stat().st_size
    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    q_old = text("""
        SELECT
            "IdRetiroLaboralAdjunto",
            "RutaArchivo"
        FROM public."RetiroLaboralAdjunto"
        WHERE "IdRetiroLaboral" = :id_retiro_laboral
          AND "IdTipoDocumentoRetiro" = :id_tipo_documento_retiro
          AND COALESCE("Eliminado", false) = false
          AND COALESCE("Activo", true) = true
        ORDER BY "IdRetiroLaboralAdjunto" DESC
        LIMIT 1;
    """)

    old = db.execute(q_old, {
        "id_retiro_laboral": id_retiro_laboral,
        "id_tipo_documento_retiro": TIPO_DOC_CARTA_FINALIZACION
    }).mappings().first()

    if old:
        q_desactivar = text("""
            UPDATE public."RetiroLaboralAdjunto"
            SET
                "Activo" = false,
                "Eliminado" = true,
                "FechaActualizacion" = now(),
                "UsuarioActualizacion" = :usuario_actualizacion
            WHERE "IdRetiroLaboralAdjunto" = :id_adjunto;
        """)
        db.execute(q_desactivar, {
            "id_adjunto": old["IdRetiroLaboralAdjunto"],
            "usuario_actualizacion": usuario_actualizacion
        })

        ruta_old = Path(old["RutaArchivo"]) if old.get("RutaArchivo") else None
        if ruta_old and ruta_old.exists():
            ruta_old.unlink(missing_ok=True)

    q_insert = text("""
        INSERT INTO public."RetiroLaboralAdjunto" (
            "IdRetiroLaboral",
            "IdTipoDocumentoRetiro",
            "NombreArchivo",
            "NombreArchivoOriginal",
            "RutaArchivo",
            "ExtensionArchivo",
            "PesoArchivo",
            "Observacion",
            "OrigenArchivo",
            "MimeType",
            "Activo",
            "Eliminado",
            "FechaCreacion",
            "FechaActualizacion",
            "CreadoPor",
            "UsuarioActualizacion"
        )
        VALUES (
            :id_retiro_laboral,
            :id_tipo_documento_retiro,
            :nombre_archivo,
            :nombre_archivo_original,
            :ruta_archivo,
            :extension_archivo,
            :peso_archivo,
            :observacion,
            'GENERADO',
            :mime_type,
            true,
            false,
            now(),
            now(),
            :creado_por,
            :usuario_actualizacion
        )
        RETURNING
            "IdRetiroLaboralAdjunto",
            "IdRetiroLaboral",
            "IdTipoDocumentoRetiro",
            "NombreArchivo",
            "NombreArchivoOriginal",
            "RutaArchivo",
            "ExtensionArchivo",
            "PesoArchivo",
            "Observacion",
            "OrigenArchivo",
            "MimeType",
            "Activo";
    """)

    row = db.execute(q_insert, {
        "id_retiro_laboral": id_retiro_laboral,
        "id_tipo_documento_retiro": TIPO_DOC_CARTA_FINALIZACION,
        "nombre_archivo": nombre_archivo,
        "nombre_archivo_original": nombre_original,
        "ruta_archivo": ruta_archivo,
        "extension_archivo": extension_archivo,
        "peso_archivo": peso_archivo,
        "observacion": "Documento generado automáticamente: Carta de finalización",
        "mime_type": mime_type,
        "creado_por": usuario_actualizacion,
        "usuario_actualizacion": usuario_actualizacion,
    }).mappings().first()

    db.commit()
    return dict(row)


def generar_y_registrar_paquete_retiro(
    db,
    id_retiro_laboral: int,
    usuario_actualizacion: str = "RRLL"
):
    output_path = generar_paquete_retiro(db, id_retiro_laboral)

    if not output_path.exists():
        raise FileNotFoundError("No se pudo generar físicamente el documento.")

    nombre_archivo = output_path.name
    nombre_original = output_path.name
    ruta_archivo = str(output_path).replace("\\", "/")
    extension_archivo = output_path.suffix.lower()
    peso_archivo = output_path.stat().st_size
    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    q_old = text("""
        SELECT
            "IdRetiroLaboralAdjunto",
            "RutaArchivo"
        FROM public."RetiroLaboralAdjunto"
        WHERE "IdRetiroLaboral" = :id_retiro_laboral
          AND "IdTipoDocumentoRetiro" = :id_tipo_documento_retiro
          AND COALESCE("Eliminado", false) = false
          AND COALESCE("Activo", true) = true
        ORDER BY "IdRetiroLaboralAdjunto" DESC
        LIMIT 1;
    """)

    old = db.execute(q_old, {
        "id_retiro_laboral": id_retiro_laboral,
        "id_tipo_documento_retiro": TIPO_DOC_PAQUETE_RETIRO
    }).mappings().first()

    if old:
        q_desactivar = text("""
            UPDATE public."RetiroLaboralAdjunto"
            SET
                "Activo" = false,
                "Eliminado" = true,
                "FechaActualizacion" = now(),
                "UsuarioActualizacion" = :usuario_actualizacion
            WHERE "IdRetiroLaboralAdjunto" = :id_adjunto;
        """)
        db.execute(q_desactivar, {
            "id_adjunto": old["IdRetiroLaboralAdjunto"],
            "usuario_actualizacion": usuario_actualizacion
        })

        ruta_old = Path(old["RutaArchivo"]) if old.get("RutaArchivo") else None
        if ruta_old and ruta_old.exists():
            ruta_old.unlink(missing_ok=True)

    q_insert = text("""
        INSERT INTO public."RetiroLaboralAdjunto" (
            "IdRetiroLaboral",
            "IdTipoDocumentoRetiro",
            "NombreArchivo",
            "NombreArchivoOriginal",
            "RutaArchivo",
            "ExtensionArchivo",
            "PesoArchivo",
            "Observacion",
            "OrigenArchivo",
            "MimeType",
            "Activo",
            "Eliminado",
            "FechaCreacion",
            "FechaActualizacion",
            "CreadoPor",
            "UsuarioActualizacion"
        )
        VALUES (
            :id_retiro_laboral,
            :id_tipo_documento_retiro,
            :nombre_archivo,
            :nombre_archivo_original,
            :ruta_archivo,
            :extension_archivo,
            :peso_archivo,
            :observacion,
            'GENERADO',
            :mime_type,
            true,
            false,
            now(),
            now(),
            :creado_por,
            :usuario_actualizacion
        )
        RETURNING
            "IdRetiroLaboralAdjunto",
            "IdRetiroLaboral",
            "IdTipoDocumentoRetiro",
            "NombreArchivo",
            "NombreArchivoOriginal",
            "RutaArchivo",
            "ExtensionArchivo",
            "PesoArchivo",
            "Observacion",
            "OrigenArchivo",
            "MimeType",
            "Activo";
    """)

    row = db.execute(q_insert, {
        "id_retiro_laboral": id_retiro_laboral,
        "id_tipo_documento_retiro": TIPO_DOC_PAQUETE_RETIRO,
        "nombre_archivo": nombre_archivo,
        "nombre_archivo_original": nombre_original,
        "ruta_archivo": ruta_archivo,
        "extension_archivo": extension_archivo,
        "peso_archivo": peso_archivo,
        "observacion": "Documento generado automáticamente: Paquete de retiro",
        "mime_type": mime_type,
        "creado_por": usuario_actualizacion,
        "usuario_actualizacion": usuario_actualizacion,
    }).mappings().first()

    db.commit()
    return dict(row)