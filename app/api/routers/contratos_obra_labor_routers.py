from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from sqlalchemy import text
from infrastructure.db.deps import get_db

from docx import Document
from io import BytesIO
from datetime import date, datetime

router = APIRouter(
    prefix="/api/contratos-obra-labor",
    tags=["contratos-obra-labor"],
)

TEMPLATE_PATH = "app/templates/CONTRATO_OBRA_LABOR.docx"  # <-- pon aquí tu plantilla

def _to_str(v):
    return "" if v is None else str(v)

def _to_date_only(v):
    # acepta date, datetime o string ISO
    if v is None:
        return None
    if isinstance(v, date) and not isinstance(v, datetime):
        return v
    if isinstance(v, datetime):
        return v.date()
    try:
        return datetime.fromisoformat(str(v)).date()
    except Exception:
        return None

def _replace_in_doc(doc: Document, mapping: dict):
    # Reemplaza en párrafos
    for p in doc.paragraphs:
        if any(k in p.text for k in mapping.keys()):
            txt = p.text
            for k, val in mapping.items():
                txt = txt.replace(k, val)
            p.text = txt

    # Reemplaza en tablas/celdas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if any(k in p.text for k in mapping.keys()):
                        txt = p.text
                        for k, val in mapping.items():
                            txt = txt.replace(k, val)
                        p.text = txt

@router.get("/{id_registro_personal}")
def descargar_contrato(id_registro_personal: int, db: Session = Depends(get_db)):
    # 1) Datos base de RegistroPersonal (ajusta columnas si en tu BD cambian)
    rp = {}
    try:
        q = text("""
            SELECT
              "Nombres", "Apellidos", "NumeroIdentificacion",
              "LugarExpedicion", "Direccion", "Barrio",
              "Telefono", "Celular", "CorreoElectronico",
              "FechaNacimiento"
            FROM "RegistroPersonal"
            WHERE "IdRegistroPersonal" = :id
            LIMIT 1
        """)
        row = db.execute(q, {"id": id_registro_personal}).mappings().first()
        if not row:
            raise HTTPException(status_code=404, detail="No existe el aspirante (RegistroPersonal)")
        rp = dict(row)
    except HTTPException:
        raise
    except Exception:
        # Si tu tabla/columnas difieren, aquí te explotaría. Mejor te muestra el error.
        raise HTTPException(status_code=500, detail="Error consultando RegistroPersonal (revisa columnas/tabla)")

    # 2) Datos de proceso (si existen) - tu endpoint ya existe: datos-proceso-aspirante
    dp = {}
    try:
        row = db.execute(
            text("""
              SELECT "FechaProceso", "IdTipoCargo", "HaTrabajadoAntesEmpresa"
              FROM "DatosProcesoAspirante"
              WHERE "IdRegistroPersonal" = :id
              LIMIT 1
            """),
            {"id": id_registro_personal},
        ).mappings().first()
        dp = dict(row) if row else {}
    except Exception:
        dp = {}

    # 3) Asignación cargo/cliente/salario (si existe)
    acc = {}
    try:
        row = db.execute(
            text("""
              SELECT "Cargo", "Salario", "Cliente"
              FROM "AsignacionCargoCliente"
              WHERE "IdRegistroPersonal" = :id
              ORDER BY "FechaCreacion" DESC
              LIMIT 1
            """),
            {"id": id_registro_personal},
        ).mappings().first()
        acc = dict(row) if row else {}
    except Exception:
        acc = {}

    # ---- Armar valores (si faltan, quedan en blanco)
    nombre_trabajador = f"{_to_str(rp.get('Nombres'))} {_to_str(rp.get('Apellidos'))}".strip()
    cedula = _to_str(rp.get("NumeroIdentificacion"))
    expedida = _to_str(rp.get("LugarExpedicion"))
    direccion = _to_str(rp.get("Direccion"))
    barrio = _to_str(rp.get("Barrio"))
    tel_fijo = _to_str(rp.get("Telefono"))
    celular = _to_str(rp.get("Celular"))
    email = _to_str(rp.get("CorreoElectronico"))

    fecha_nac = _to_date_only(rp.get("FechaNacimiento"))
    fecha_nac_str = fecha_nac.strftime("%d/%m/%Y") if fecha_nac else ""

    # si ya tienes fecha de inicio labores en otra tabla, aquí la conectas
    inicio = _to_date_only(dp.get("FechaProceso"))
    inicio_dia = f"{inicio.day:02d}" if inicio else ""
    inicio_mes = f"{inicio.month:02d}" if inicio else ""
    inicio_ano = f"{inicio.year}" if inicio else ""

    cargo = _to_str(acc.get("Cargo"))
    salario = _to_str(acc.get("Salario"))

    mapping = {
        "{{NOMBRE_TRABAJADOR}}": nombre_trabajador,
        "{{CEDULA}}": cedula,
        "{{CIUDAD_EXPEDICION}}": expedida,
        "{{DIRECCION_TRABAJADOR}}": direccion,
        "{{BARRIO}}": barrio,
        "{{TELEFONO_FIJO}}": tel_fijo,
        "{{CELULAR}}": celular,
        "{{EMAIL}}": email,
        "{{FECHA_NACIMIENTO}}": fecha_nac_str,
        "{{CARGO}}": cargo,
        "{{SALARIO}}": salario,
        "{{INICIO_DIA}}": inicio_dia,
        "{{INICIO_MES}}": inicio_mes,
        "{{INICIO_ANO}}": inicio_ano,
    }

    # 4) Cargar plantilla, reemplazar y devolver
    try:
        doc = Document(TEMPLATE_PATH)
    except Exception:
        raise HTTPException(status_code=500, detail="No se pudo abrir la plantilla .docx (revisa TEMPLATE_PATH)")

    _replace_in_doc(doc, mapping)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"CONTRATO_OBRA_LABOR_{cedula or id_registro_personal}.docx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
